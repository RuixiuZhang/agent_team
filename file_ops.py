"""
文件操作 — 路径安全校验 / workspace_files 操作 / files/ 本地操作 / 文档读写
"""

import os
from typing import Dict

# 项目根目录（team 包的上一级）
_PROJECT_ROOT   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FILE_WORKSPACE  = os.path.join(_PROJECT_ROOT, "workspace_files")
LOCAL_FILE_DIR  = os.path.join(_PROJECT_ROOT, "files")

DOC_MAX_READ_CHARS = 8000


# ─────────────────────────────────────────────────────────────────────────────
# 路径安全
# ─────────────────────────────────────────────────────────────────────────────

def safe_resolve_path(relative_path: str) -> str:
    """解析相对路径到 workspace_files/ 目录，防止路径越界"""
    cleaned    = relative_path.strip().strip('"').strip("'")
    if os.path.isabs(cleaned):
        raise ValueError(f"不允许绝对路径: {cleaned}")
    normalized = os.path.normpath(cleaned)
    if normalized.startswith("..") or "/../" in normalized or normalized == "..":
        raise ValueError(f"路径越界: {cleaned}")
    full_path  = os.path.join(FILE_WORKSPACE, normalized)
    real_base  = os.path.realpath(FILE_WORKSPACE)
    real_target = os.path.realpath(full_path)
    if not real_target.startswith(real_base + os.sep) and real_target != real_base:
        raise ValueError(f"路径越界: {cleaned}")
    return full_path


def safe_resolve_local_path(relative_path: str) -> str:
    """解析相对路径到 files/ 目录，防止路径越界"""
    cleaned = relative_path.strip().strip('"').strip("'")
    if os.path.isabs(cleaned):
        raise ValueError(f"不允许绝对路径: {cleaned}")
    normalized = os.path.normpath(cleaned)
    if normalized.startswith("..") or "/../" in normalized or normalized == "..":
        raise ValueError(f"路径越界: {cleaned}")
    full_path = os.path.join(LOCAL_FILE_DIR, normalized)
    real_base = os.path.realpath(LOCAL_FILE_DIR)
    real_target = os.path.realpath(full_path)
    if not real_target.startswith(real_base + os.sep) and real_target != real_base:
        raise ValueError(f"路径越界: {cleaned}")
    return full_path


# ─────────────────────────────────────────────────────────────────────────────
# workspace_files 文件操作
# ─────────────────────────────────────────────────────────────────────────────

def execute_file_op(op: dict) -> str:
    action   = op.get("op", "").lower()
    rel_path = op.get("path", "")

    if action == "read":
        full = safe_resolve_path(rel_path)
        if not os.path.exists(full):
            return f"[read] 错误：文件不存在 '{rel_path}'"
        if not os.path.isfile(full):
            return f"[read] 错误：路径不是文件 '{rel_path}'"
        size = os.path.getsize(full)
        if size > 512 * 1024:
            return f"[read] 错误：文件过大 ({size} bytes)，限制 512KB"
        with open(full, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return f"[read] 已读取 '{rel_path}' ({len(content)} 字):\n{content}"

    elif action == "write":
        full    = safe_resolve_path(rel_path)
        content = op.get("content", "")
        parent  = os.path.dirname(full)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(full, "w", encoding="utf-8") as f:
            f.write(content)
        return f"[write] 已写入 '{rel_path}' ({len(content)} 字)"

    elif action == "append":
        full    = safe_resolve_path(rel_path)
        content = op.get("content", "")
        parent  = os.path.dirname(full)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(full, "a", encoding="utf-8") as f:
            f.write(content)
        return f"[append] 已追加 '{rel_path}' ({len(content)} 字)"

    elif action == "list":
        full = (
            safe_resolve_path(rel_path)
            if rel_path and rel_path != "."
            else FILE_WORKSPACE
        )
        if not os.path.isdir(full):
            return f"[list] 错误：目录不存在 '{rel_path}'"
        entries = []
        for name in sorted(os.listdir(full)):
            entry_path = os.path.join(full, name)
            if os.path.isdir(entry_path):
                entries.append(f"  📁 {name}/")
            else:
                entries.append(f"  📄 {name} ({os.path.getsize(entry_path)} bytes)")
        if not entries:
            return f"[list] 目录 '{rel_path}' 为空"
        return f"[list] 目录 '{rel_path}':\n" + "\n".join(entries)

    else:
        return f"[error] 未知操作: '{action}' (支持: read/write/append/list)"


# ─────────────────────────────────────────────────────────────────────────────
# files/ 本地文件操作
# ─────────────────────────────────────────────────────────────────────────────

def execute_local_file_op(op: dict) -> str:
    action = op.get("op", "").lower()
    rel_path = op.get("path", "")

    if action == "read":
        full = safe_resolve_local_path(rel_path)
        if not os.path.exists(full):
            return f"[read] 错误：文件不存在 '{rel_path}'"
        if not os.path.isfile(full):
            return f"[read] 错误：路径不是文件 '{rel_path}'"
        size = os.path.getsize(full)
        if size > 512 * 1024:
            return f"[read] 错误：文件过大 ({size} bytes)，限制 512KB"
        with open(full, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return f"[read] 已读取 '{rel_path}' ({len(content)} 字):\n{content}"

    elif action == "write":
        full = safe_resolve_local_path(rel_path)
        content = op.get("content", "")
        parent = os.path.dirname(full)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(full, "w", encoding="utf-8") as f:
            f.write(content)
        return f"[write] 已写入 files/{rel_path} ({len(content)} 字)"

    elif action == "append":
        full = safe_resolve_local_path(rel_path)
        content = op.get("content", "")
        parent = os.path.dirname(full)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(full, "a", encoding="utf-8") as f:
            f.write(content)
        return f"[append] 已追加 files/{rel_path} ({len(content)} 字)"

    elif action == "mkdir":
        full = safe_resolve_local_path(rel_path)
        os.makedirs(full, exist_ok=True)
        return f"[mkdir] 已创建目录 files/{rel_path}"

    elif action == "list":
        full = (
            safe_resolve_local_path(rel_path)
            if rel_path and rel_path != "."
            else LOCAL_FILE_DIR
        )
        if not os.path.isdir(full):
            return f"[list] 错误：目录不存在 '{rel_path}'"
        entries = []
        for name in sorted(os.listdir(full)):
            entry_path = os.path.join(full, name)
            if os.path.isdir(entry_path):
                entries.append(f"  📁 {name}/")
            else:
                entries.append(f"  📄 {name} ({os.path.getsize(entry_path)} bytes)")
        if not entries:
            return f"[list] 目录 files/{rel_path} 为空"
        return f"[list] 目录 files/{rel_path}:\n" + "\n".join(entries)

    else:
        return f"[error] 未知操作: '{action}' (支持: read/write/append/mkdir/list)"


# ─────────────────────────────────────────────────────────────────────────────
# 文档读写
# ─────────────────────────────────────────────────────────────────────────────

def read_document(file_path: str) -> str:
    """读取各种格式的文档，返回文本内容"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        import pdfplumber
        texts = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t:
                    texts.append(f"--- 第{i+1}页 ---\n{t}")
        return "\n".join(texts)[:DOC_MAX_READ_CHARS]

    elif ext == ".docx":
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)[:DOC_MAX_READ_CHARS]

    elif ext == ".xlsx":
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                parts.append("\t".join(str(c) if c is not None else "" for c in row))
        wb.close()
        return "\n".join(parts)[:DOC_MAX_READ_CHARS]

    elif ext == ".csv":
        import csv as csv_mod
        rows = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv_mod.reader(f)
            for i, row in enumerate(reader):
                rows.append("\t".join(row))
                if i > 500:
                    rows.append("... (已截断)")
                    break
        return "\n".join(rows)[:DOC_MAX_READ_CHARS]

    elif ext in (".txt", ".md", ".json", ".log"):
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read(DOC_MAX_READ_CHARS)

    else:
        return f"不支持的文件格式: {ext}"


def write_docx(file_path: str, title: str, paragraphs: list) -> str:
    """生成 Word 文档"""
    from docx import Document
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for p in paragraphs:
        if isinstance(p, str):
            doc.add_paragraph(p)
        elif isinstance(p, dict):
            text = p.get("text", "")
            level = p.get("heading_level", 0)
            if level:
                doc.add_heading(text, level=min(level, 4))
            else:
                doc.add_paragraph(text)
    parent = os.path.dirname(file_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    doc.save(file_path)
    return f"[write_docx] 已生成 Word 文档 ({len(paragraphs)} 段)"


def write_xlsx(file_path: str, sheets: dict) -> str:
    """生成 Excel 文档"""
    import openpyxl
    wb = openpyxl.Workbook()
    first = True
    for sheet_name, data in sheets.items():
        if first:
            ws = wb.active
            ws.title = sheet_name
            first = False
        else:
            ws = wb.create_sheet(sheet_name)
        headers = data.get("headers", [])
        rows = data.get("rows", [])
        if headers:
            ws.append(headers)
        for row in rows:
            ws.append(row)
    parent = os.path.dirname(file_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    wb.save(file_path)
    total_rows = sum(len(d.get("rows", [])) for d in sheets.values())
    return f"[write_xlsx] 已生成 Excel ({len(sheets)} 个 Sheet, {total_rows} 行数据)"


def write_csv_file(file_path: str, headers: list, rows: list) -> str:
    """生成 CSV 文件"""
    import csv as csv_mod
    parent = os.path.dirname(file_path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    with open(file_path, "w", encoding="utf-8", newline="") as f:
        writer = csv_mod.writer(f)
        if headers:
            writer.writerow(headers)
        writer.writerows(rows)
    return f"[write_csv] 已生成 CSV ({len(rows)} 行)"


def execute_doc_op(op: dict) -> str:
    """执行单个文档操作"""
    action = op.get("op", "").lower()
    rel_path = op.get("path", "")

    if action == "read_doc":
        full = safe_resolve_local_path(rel_path)
        if not os.path.exists(full):
            return f"[read_doc] 错误：文件不存在 '{rel_path}'"
        if not os.path.isfile(full):
            return f"[read_doc] 错误：路径不是文件 '{rel_path}'"
        size = os.path.getsize(full)
        if size > 50 * 1024 * 1024:
            return f"[read_doc] 错误：文件过大 ({size} bytes)，限制 50MB"
        try:
            content = read_document(full)
            return f"[read_doc] 已读取 '{rel_path}' ({len(content)} 字):\n{content}"
        except Exception as e:
            return f"[read_doc] 读取失败: {type(e).__name__}: {e}"

    elif action == "write_docx":
        full = safe_resolve_local_path(rel_path)
        title = op.get("title", "")
        paragraphs = op.get("paragraphs", [])
        try:
            return write_docx(full, title, paragraphs)
        except Exception as e:
            return f"[write_docx] 生成失败: {type(e).__name__}: {e}"

    elif action == "write_xlsx":
        full = safe_resolve_local_path(rel_path)
        sheets = op.get("sheets", {})
        try:
            return write_xlsx(full, sheets)
        except Exception as e:
            return f"[write_xlsx] 生成失败: {type(e).__name__}: {e}"

    elif action == "write_csv":
        full = safe_resolve_local_path(rel_path)
        headers = op.get("headers", [])
        rows = op.get("rows", [])
        try:
            return write_csv_file(full, headers, rows)
        except Exception as e:
            return f"[write_csv] 生成失败: {type(e).__name__}: {e}"

    elif action == "write_text":
        full = safe_resolve_local_path(rel_path)
        content = op.get("content", "")
        parent = os.path.dirname(full)
        if parent:
            os.makedirs(parent, exist_ok=True)
        with open(full, "w", encoding="utf-8") as f:
            f.write(content)
        return f"[write_text] 已写入 '{rel_path}' ({len(content)} 字)"

    elif action == "list_docs":
        full = (
            safe_resolve_local_path(rel_path)
            if rel_path and rel_path != "."
            else LOCAL_FILE_DIR
        )
        if not os.path.isdir(full):
            return f"[list_docs] 错误：目录不存在 '{rel_path}'"
        doc_exts = {".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md", ".json", ".log"}
        entries = []
        for name in sorted(os.listdir(full)):
            entry_path = os.path.join(full, name)
            if os.path.isdir(entry_path):
                entries.append(f"  📁 {name}/")
            else:
                ext = os.path.splitext(name)[1].lower()
                tag = " [文档]" if ext in doc_exts else ""
                entries.append(f"  📄 {name} ({os.path.getsize(entry_path)} bytes){tag}")
        if not entries:
            return f"[list_docs] 目录 '{rel_path}' 为空"
        return f"[list_docs] 目录 '{rel_path}':\n" + "\n".join(entries)

    else:
        return f"[error] 未知操作: '{action}' (支持: read_doc/write_docx/write_xlsx/write_csv/write_text/list_docs)"
